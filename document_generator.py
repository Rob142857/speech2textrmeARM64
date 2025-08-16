#!/usr/bin/env python3
"""
Document Generation Module for ARM64 Whisper Transcription
Creates formatted DOCX documents with proper structure and metadata
"""

import os
import sys
from pathlib import Path
from typing import Dict, Any, Optional
import time
from datetime import datetime

try:
    from docx import Document
    from docx.shared import Inches, Pt
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.style import WD_STYLE_TYPE
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False
    print("⚠️  python-docx not available - DOCX generation will be limited")

class DocumentGenerator:
    """
    Professional document generator for transcription results.
    Creates structured DOCX files with proper formatting and metadata.
    """
    
    def __init__(self):
        """Initialize document generator."""
        self.docx_available = DOCX_AVAILABLE
        
        if self.docx_available:
            print("📄 Document generator ready (DOCX support enabled)")
        else:
            print("📄 Document generator ready (text-only mode)")
    
    def create_document(
        self,
        transcription_text: str,
        output_path: str,
        title: str = "Speech Transcription",
        metadata: Optional[Dict[str, Any]] = None
    ) -> str:
        """
        Create a formatted document from transcription text.
        
        Args:
            transcription_text: The transcribed text
            output_path: Path for output document
            title: Document title
            metadata: Additional metadata to include
            
        Returns:
            Path to created document
        """
        output_path = Path(output_path)
        
        if self.docx_available and output_path.suffix.lower() == '.docx':
            return self._create_docx_document(transcription_text, str(output_path), title, metadata)
        else:
            return self._create_text_document(transcription_text, str(output_path), title, metadata)
    
    def _create_docx_document(
        self,
        transcription_text: str,
        output_path: str,
        title: str,
        metadata: Optional[Dict[str, Any]] = None
    ) -> str:
        """Create a professionally formatted DOCX document."""
        try:
            # Create document
            doc = Document()
            
            # Set document properties
            properties = doc.core_properties
            properties.title = title
            properties.author = "ARM64 Whisper Transcriber"
            properties.subject = "Speech-to-Text Transcription"
            properties.created = datetime.now()
            
            # Add title
            title_paragraph = doc.add_heading(title, level=0)
            title_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            # Add subtitle with timestamp
            subtitle = doc.add_paragraph()
            subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
            subtitle_run = subtitle.add_run(f"Generated on {datetime.now().strftime('%B %d, %Y at %I:%M %p')}")
            subtitle_run.font.size = Pt(12)
            subtitle_run.italic = True
            
            # Add separator
            doc.add_paragraph("_" * 80).alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            # Add metadata section if provided
            if metadata:
                self._add_metadata_section(doc, metadata)
            
            # Add transcription content
            self._add_transcription_content(doc, transcription_text)
            
            # Add footer
            self._add_document_footer(doc)
            
            # Save document
            doc.save(output_path)
            print(f"✅ DOCX document created: {Path(output_path).name}")
            
            return output_path
            
        except Exception as e:
            print(f"❌ DOCX creation failed: {e}")
            # Fallback to text document
            return self._create_text_document(transcription_text, output_path.replace('.docx', '.txt'), title, metadata)
    
    def _add_metadata_section(self, doc: Document, metadata: Dict[str, Any]) -> None:
        """Add metadata section to document."""
        # Metadata heading
        doc.add_heading("Transcription Information", level=1)
        
        # Create metadata table
        table = doc.add_table(rows=0, cols=2)
        table.style = 'Table Grid'
        
        # Add metadata rows
        for key, value in metadata.items():
            if value is not None:
                # Format key
                display_key = key.replace('_', ' ').title()
                
                # Format value
                if isinstance(value, bool):
                    display_value = "Yes" if value else "No"
                elif isinstance(value, (int, float)) and key.endswith('_seconds'):
                    display_value = f"{value:.1f} seconds"
                elif isinstance(value, dict):
                    display_value = str(value)
                else:
                    display_value = str(value)
                
                # Add row
                row = table.add_row()
                row.cells[0].text = display_key
                row.cells[1].text = display_value
                
                # Format cells
                row.cells[0].paragraphs[0].runs[0].bold = True
        
        doc.add_paragraph()  # Add space after table
    
    def _add_transcription_content(self, doc: Document, transcription_text: str) -> None:
        """Add formatted transcription content."""
        # Content heading
        doc.add_heading("Transcription", level=1)
        
        if not transcription_text or not transcription_text.strip():
            # No content case
            no_content = doc.add_paragraph("No transcription content available.")
            no_content.runs[0].italic = True
            return
        
        # Split text into paragraphs
        paragraphs = self._format_into_paragraphs(transcription_text)
        
        for paragraph_text in paragraphs:
            if paragraph_text.strip():
                para = doc.add_paragraph(paragraph_text.strip())
                
                # Set paragraph formatting
                paragraph_format = para.paragraph_format
                paragraph_format.space_after = Pt(6)
                paragraph_format.first_line_indent = Inches(0.5)
                paragraph_format.line_spacing = 1.15
    
    def _format_into_paragraphs(self, text: str) -> list:
        """Format text into logical paragraphs."""
        if not text:
            return []
        
        # Basic paragraph detection
        # This is a simple heuristic - can be enhanced
        sentences = text.split('. ')
        paragraphs = []
        current_paragraph = []
        
        for sentence in sentences:
            sentence = sentence.strip()
            if not sentence:
                continue
            
            # Ensure sentence ends properly
            if not sentence.endswith('.') and sentence != sentences[-1]:
                sentence += '.'
            
            current_paragraph.append(sentence)
            
            # Start new paragraph on certain cues
            if (len(current_paragraph) >= 3 or
                any(sentence.lower().startswith(cue) for cue in 
                    ['so ', 'now ', 'then ', 'and then ', 'but ', 'however ', 'meanwhile '])):
                
                if len(current_paragraph) > 0:
                    paragraphs.append(' '.join(current_paragraph))
                    current_paragraph = []
        
        # Add remaining content
        if current_paragraph:
            paragraphs.append(' '.join(current_paragraph))
        
        return paragraphs
    
    def _add_document_footer(self, doc: Document) -> None:
        """Add footer information."""
        doc.add_page_break()
        
        # Footer heading
        doc.add_heading("Technical Information", level=2)
        
        footer_text = """
        This transcription was generated using ARM64 Whisper Transcription Engine.
        
        Key Features:
        • Local processing (no cloud dependencies)
        • NPU acceleration on compatible ARM64 devices
        • Multi-format audio/video support
        • Automatic punctuation and paragraphing
        
        For questions or support, please refer to the project documentation.
        """
        
        footer_para = doc.add_paragraph(footer_text.strip())
        footer_para.runs[0].font.size = Pt(10)
        footer_para.runs[0].italic = True
    
    def _create_text_document(
        self,
        transcription_text: str,
        output_path: str,
        title: str,
        metadata: Optional[Dict[str, Any]] = None
    ) -> str:
        """Create a formatted text document as fallback."""
        try:
            output_path = str(output_path)
            if not output_path.endswith('.txt'):
                output_path = output_path.rsplit('.', 1)[0] + '.txt'
            
            with open(output_path, 'w', encoding='utf-8') as f:
                # Write header
                f.write(f"{title}\n")
                f.write("=" * len(title) + "\n\n")
                
                # Write timestamp
                f.write(f"Generated on {datetime.now().strftime('%B %d, %Y at %I:%M %p')}\n")
                f.write("-" * 50 + "\n\n")
                
                # Write metadata
                if metadata:
                    f.write("TRANSCRIPTION INFORMATION\n")
                    f.write("-" * 25 + "\n")
                    
                    for key, value in metadata.items():
                        if value is not None:
                            display_key = key.replace('_', ' ').title()
                            
                            if isinstance(value, bool):
                                display_value = "Yes" if value else "No"
                            elif isinstance(value, (int, float)) and key.endswith('_seconds'):
                                display_value = f"{value:.1f} seconds"
                            else:
                                display_value = str(value)
                            
                            f.write(f"{display_key}: {display_value}\n")
                    
                    f.write("\n" + "-" * 50 + "\n\n")
                
                # Write transcription
                f.write("TRANSCRIPTION\n")
                f.write("-" * 13 + "\n\n")
                
                if transcription_text and transcription_text.strip():
                    # Format into paragraphs
                    paragraphs = self._format_into_paragraphs(transcription_text)
                    for paragraph in paragraphs:
                        if paragraph.strip():
                            f.write(f"{paragraph.strip()}\n\n")
                else:
                    f.write("No transcription content available.\n\n")
                
                # Write footer
                f.write("-" * 50 + "\n")
                f.write("Generated by ARM64 Whisper Transcription Engine\n")
                f.write("Local processing with NPU acceleration\n")
            
            print(f"✅ Text document created: {Path(output_path).name}")
            return output_path
            
        except Exception as e:
            print(f"❌ Text document creation failed: {e}")
            raise
    
    def create_batch_report(
        self,
        results: Dict[str, Dict[str, Any]],
        output_path: str,
        title: str = "Batch Transcription Report"
    ) -> str:
        """
        Create a summary report for batch transcription results.
        
        Args:
            results: Dictionary of transcription results
            output_path: Path for output report
            title: Report title
            
        Returns:
            Path to created report
        """
        try:
            summary_data = {
                "total_files": len(results),
                "successful": sum(1 for r in results.values() if r.get("success", False)),
                "failed": sum(1 for r in results.values() if not r.get("success", False)),
                "total_processing_time": sum(r.get("processing_time", 0) for r in results.values()),
                "average_processing_time": 0,
                "report_generated": datetime.now().strftime('%Y-%m-%d %H:%M:%S')
            }
            
            if summary_data["successful"] > 0:
                summary_data["average_processing_time"] = (
                    summary_data["total_processing_time"] / summary_data["successful"]
                )
            
            if self.docx_available and output_path.endswith('.docx'):
                return self._create_batch_docx_report(results, output_path, title, summary_data)
            else:
                return self._create_batch_text_report(results, output_path, title, summary_data)
                
        except Exception as e:
            print(f"❌ Batch report creation failed: {e}")
            raise
    
    def _create_batch_docx_report(
        self,
        results: Dict[str, Dict[str, Any]],
        output_path: str,
        title: str,
        summary_data: Dict[str, Any]
    ) -> str:
        """Create DOCX batch report."""
        doc = Document()
        
        # Title
        doc.add_heading(title, level=0).alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Summary section
        doc.add_heading("Summary", level=1)
        
        summary_table = doc.add_table(rows=0, cols=2)
        summary_table.style = 'Table Grid'
        
        for key, value in summary_data.items():
            if key != "report_generated":
                row = summary_table.add_row()
                display_key = key.replace('_', ' ').title()
                
                if isinstance(value, float) and key.endswith('_time'):
                    display_value = f"{value:.1f} seconds"
                else:
                    display_value = str(value)
                
                row.cells[0].text = display_key
                row.cells[1].text = display_value
                row.cells[0].paragraphs[0].runs[0].bold = True
        
        # Results section
        doc.add_heading("Individual Results", level=1)
        
        for file_path, result in results.items():
            file_name = Path(file_path).name
            doc.add_heading(file_name, level=2)
            
            if result.get("success", False):
                doc.add_paragraph(f"✅ Success - Processed in {result.get('processing_time', 0):.1f} seconds")
                if result.get("transcription"):
                    preview = result["transcription"][:200] + "..." if len(result["transcription"]) > 200 else result["transcription"]
                    doc.add_paragraph(f"Preview: {preview}")
            else:
                doc.add_paragraph(f"❌ Failed - {result.get('error', 'Unknown error')}")
        
        doc.save(output_path)
        return output_path
    
    def _create_batch_text_report(
        self,
        results: Dict[str, Dict[str, Any]],
        output_path: str,
        title: str,
        summary_data: Dict[str, Any]
    ) -> str:
        """Create text batch report."""
        with open(output_path, 'w', encoding='utf-8') as f:
            f.write(f"{title}\n")
            f.write("=" * len(title) + "\n\n")
            
            f.write("SUMMARY\n")
            f.write("-------\n")
            for key, value in summary_data.items():
                display_key = key.replace('_', ' ').title()
                if isinstance(value, float) and key.endswith('_time'):
                    display_value = f"{value:.1f} seconds"
                else:
                    display_value = str(value)
                f.write(f"{display_key}: {display_value}\n")
            
            f.write("\nINDIVIDUAL RESULTS\n")
            f.write("------------------\n")
            
            for file_path, result in results.items():
                file_name = Path(file_path).name
                f.write(f"\n{file_name}:\n")
                
                if result.get("success", False):
                    f.write(f"  ✅ Success - {result.get('processing_time', 0):.1f}s\n")
                    if result.get("transcription"):
                        preview = result["transcription"][:100] + "..." if len(result["transcription"]) > 100 else result["transcription"]
                        f.write(f"  Preview: {preview}\n")
                else:
                    f.write(f"  ❌ Failed - {result.get('error', 'Unknown error')}\n")
        
        return output_path

def main():
    """Test document generator."""
    generator = DocumentGenerator()
    
    test_text = """
    This is a test transcription. It contains multiple sentences that should be formatted properly.
    
    The document generator should create professional-looking documents with proper structure.
    This includes headers, metadata, and well-formatted content.
    
    The system supports both DOCX and text output formats depending on available libraries.
    """
    
    test_metadata = {
        "source_file": "test_audio.mp3",
        "model": "medium",
        "processing_time": 45.2,
        "npu_used": True,
        "language": "en"
    }
    
    # Test DOCX creation
    if generator.docx_available:
        docx_path = "test_transcription.docx"
        generator.create_document(test_text, docx_path, "Test Transcription", test_metadata)
        print(f"Created: {docx_path}")
    
    # Test text creation
    txt_path = "test_transcription.txt"
    generator.create_document(test_text, txt_path, "Test Transcription", test_metadata)
    print(f"Created: {txt_path}")

if __name__ == "__main__":
    main()
